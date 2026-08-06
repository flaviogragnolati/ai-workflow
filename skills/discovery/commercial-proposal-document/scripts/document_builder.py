from __future__ import annotations

import json
import tempfile
import zlib
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


COLORS = {
    "mint": "69BC9B",
    "navy": "27367E",
    "black": "1C1C1C",
    "text": "1F2937",
    "muted": "6B7280",
    "line": "D8DEE9",
    "navy_tint": "EEF1FA",
    "mint_tint": "EAF7F2",
    "warning_tint": "FFF7D6",
    "white": "FFFFFF",
}

A4_WIDTH_MM = 210
A4_HEIGHT_MM = 297
BODY_LEFT_MM = 18
BODY_RIGHT_MM = 18
BODY_TOP_MM = 30
BODY_BOTTOM_MM = 29
CONTENT_WIDTH_MM = A4_WIDTH_MM - BODY_LEFT_MM - BODY_RIGHT_MM
CONTACT_EMAIL = "info@quasartech.xyz"
WEBSITE = "quasartech.xyz"
LEGAL_ENTITY = "Ingeniería Quasar SRL"


def tr(data: dict[str, Any], key: str, default: str) -> str:
    return str(data.get('labels', {}).get(key, default))


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.lstrip("#")
    return RGBColor.from_string(value)


def set_run_font(
    run,
    name: str = "Arial",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_style_font(style, name: str):
    style.font.name = name
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = COLORS["line"], size: int = 5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_mm: list[float]):
    widths_twips = [round(mm * 56.6929133858) for mm in widths_mm]
    total = sum(widths_twips)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_twips[index]
            tcpr = cell._tc.get_or_add_tcPr()
            tc_w = tcpr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tcpr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = trpr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        trpr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    cant_split = trpr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_field(run, instruction: str, fallback: str = "1"):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_page_fields(paragraph, data: dict[str, Any]):
    run = paragraph.add_run(tr(data, 'page', 'Página') + ' ')
    set_run_font(run, size=7.2, color=COLORS["muted"])
    page_run = paragraph.add_run()
    set_run_font(page_run, size=7.2, color=COLORS["muted"])
    add_field(page_run, "PAGE")
    run = paragraph.add_run(' ' + tr(data, 'of', 'de') + ' ')
    set_run_font(run, size=7.2, color=COLORS["muted"])
    pages_run = paragraph.add_run()
    set_run_font(pages_run, size=7.2, color=COLORS["muted"])
    add_field(pages_run, "NUMPAGES")


def add_text_content_control(
    paragraph,
    tag: str,
    title: str,
    display_text: str = "Haga clic o toque aquí para completar",
    size: float = 8.5,
):
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")

    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), title)
    sdt_pr.append(alias)

    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    sdt_pr.append(tag_element)

    control_id = OxmlElement("w:id")
    control_id.set(qn("w:val"), str(zlib.crc32(tag.encode("utf-8")) & 0x7FFFFFFF))
    sdt_pr.append(control_id)

    text_control = OxmlElement("w:text")
    sdt_pr.append(text_control)

    appearance = OxmlElement("w:appearance")
    appearance.set(qn("w:val"), "boundingBox")
    sdt_pr.append(appearance)

    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:cs"), "Arial")
    run_properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLORS["muted"])
    run_properties.append(color)
    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(round(size * 2)))
    run_properties.append(size_element)
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = display_text
    run.append(text)
    sdt_content.append(run)

    sdt.append(sdt_pr)
    sdt.append(sdt_content)
    paragraph._p.append(sdt)
    return sdt
def set_update_fields(document: Document):
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_styles(document: Document):
    normal = document.styles["Normal"]
    set_style_font(normal, "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(COLORS["text"])
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.08

    title = document.styles["Title"]
    set_style_font(title, "Arial")
    title.font.size = Pt(21)
    title.font.bold = True
    title.font.color.rgb = rgb(COLORS["navy"])
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.styles["Subtitle"]
    set_style_font(subtitle, "Arial")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = rgb(COLORS["muted"])
    subtitle.paragraph_format.space_after = Pt(8)

    h1 = document.styles["Heading 1"]
    set_style_font(h1, "Arial")
    h1.font.size = Pt(15.5)
    h1.font.bold = True
    h1.font.color.rgb = rgb(COLORS["navy"])
    h1.paragraph_format.space_before = Pt(11)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True

    h2 = document.styles["Heading 2"]
    set_style_font(h2, "Arial")
    h2.font.size = Pt(11.5)
    h2.font.bold = True
    h2.font.color.rgb = rgb(COLORS["navy"])
    h2.paragraph_format.space_before = Pt(7)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True

    h3 = document.styles["Heading 3"]
    set_style_font(h3, "Arial")
    h3.font.size = Pt(10)
    h3.font.bold = True
    h3.font.color.rgb = rgb(COLORS["text"])
    h3.paragraph_format.space_before = Pt(5)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = document.styles[list_name]
        set_style_font(style, "Arial")
        style.font.size = Pt(9.25)
        style.font.color.rgb = rgb(COLORS["text"])
        style.paragraph_format.space_after = Pt(2.5)

    for name, size, color, bold in (
        ("Part Label", 8.5, COLORS["mint"], True),
        ("Section Title", 21, COLORS["navy"], True),
        ("Small Text", 7.5, COLORS["muted"], False),
        ("Lead", 10.5, COLORS["text"], False),
    ):
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = bold
        style.paragraph_format.space_after = Pt(4)
        if name in {"Part Label", "Section Title"}:
            style.paragraph_format.keep_with_next = True


def load_font(path: Path, size: int):
    try:
        return ImageFont.truetype(str(path), size=size)
    except Exception:
        return ImageFont.truetype("arial.ttf", size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = str(text).split()
    if not words:
        return [""]
    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font,
    fill: str,
    max_width: int,
    line_gap: int = 8,
    max_lines: int | None = None,
):
    x, y = xy
    lines = wrap_text(draw, text, font, max_width)
    if max_lines is not None:
        lines = lines[:max_lines]
    bbox = draw.textbbox((x, y), "Ag", font=font)
    line_height = bbox[3] - bbox[1]
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height + line_gap
    return y


def build_cover_image(data: dict[str, Any], assets: Path, output: Path):
    metadata = data["metadata"]
    width, height = 1240, 1754
    canvas = Image.new("RGB", (width, height), "white")
    pattern = Image.open(assets / "cover-pattern.png").convert("RGB")
    pattern_height = round(width * pattern.height / pattern.width)
    pattern = pattern.resize((width, pattern_height))
    canvas.paste(pattern, (0, 0))
    draw = ImageDraw.Draw(canvas)
    draw.rectangle((0, pattern_height - 5, width, pattern_height + 4), fill="#69BC9B")

    extra = assets / "SuperaGothic-ExtraBold.otf"
    regular = assets / "SuperaGothic-Regular.otf"
    font_label = load_font(extra, 24)
    font_client = load_font(extra, 39)
    font_project = load_font(extra, 34)
    font_meta = load_font(regular, 20)
    font_contact = load_font(regular, 21)

    x = 110
    y = pattern_height + 105
    draw.text((x, y), tr(data, 'cover_label', 'PROPUESTA COMERCIAL'), font=font_label, fill="#69BC9B")
    y += 48
    y = draw_wrapped(
        draw,
        (x, y),
        metadata["client"],
        font_client,
        "#69BC9B",
        900,
        line_gap=5,
        max_lines=2,
    )
    y += 18
    y = draw_wrapped(
        draw,
        (x, y),
        metadata["project_name"],
        font_project,
        "#27367E",
        900,
        line_gap=5,
        max_lines=3,
    )
    subtitle = metadata.get("project_subtitle", "")
    if subtitle:
        y += 12
        draw_wrapped(draw, (x, y), subtitle, font_meta, "#6B7280", 900, line_gap=4, max_lines=2)

    meta_y = 1480
    meta_text = f"{tr(data, 'version', 'Versión')} {metadata['version']}  |  {metadata['issue_date']}"
    draw.text((x, meta_y), meta_text, font=font_meta, fill="#6B7280")

    contact = f"{CONTACT_EMAIL}  |  {WEBSITE}"
    draw.text((x, 1620), contact, font=font_contact, fill="#27367E")

    logo = Image.open(assets / "logo-horizontal-color.png").convert("RGBA")
    target_width = 470
    target_height = round(logo.height * target_width / logo.width)
    logo = logo.resize((target_width, target_height))
    canvas.paste(logo, (width - target_width - 95, 1570), logo)
    canvas.save(output, quality=95)


def build_footer_band(data: dict[str, Any], assets: Path, output: Path):
    width, height = 1240, 118
    pattern = Image.open(assets / "cover-pattern.png").convert("RGB")
    crop = pattern.crop((0, 0, pattern.width, min(150, pattern.height)))
    band = crop.resize((width, height))
    overlay = Image.new("RGBA", band.size, (20, 27, 66, 65))
    band = Image.alpha_composite(band.convert("RGBA"), overlay)
    draw = ImageDraw.Draw(band)
    font = load_font(assets / "SuperaGothic-Regular.otf", 18)
    contact = f"{CONTACT_EMAIL}  |  {WEBSITE}"
    bbox = draw.textbbox((0, 0), contact, font=font)
    x = (width - (bbox[2] - bbox[0])) // 2
    y = (height - (bbox[3] - bbox[1])) // 2 - 2
    draw.text((x, y), contact, font=font, fill="white")
    band.convert("RGB").save(output, quality=92)


def configure_cover_section(section):
    section.page_width = Mm(A4_WIDTH_MM)
    section.page_height = Mm(A4_HEIGHT_MM)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)


def configure_body_section(section, assets: Path, footer_band: Path, data: dict[str, Any]):
    metadata = data['metadata']
    section.page_width = Mm(A4_WIDTH_MM)
    section.page_height = Mm(A4_HEIGHT_MM)
    section.top_margin = Mm(BODY_TOP_MM)
    section.bottom_margin = Mm(BODY_BOTTOM_MM)
    section.left_margin = Mm(BODY_LEFT_MM)
    section.right_margin = Mm(BODY_RIGHT_MM)
    section.header_distance = Mm(0)
    section.footer_distance = Mm(0)
    section.different_first_page_header_footer = False

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.paragraph_format.left_indent = Mm(-BODY_LEFT_MM)
    paragraph.paragraph_format.right_indent = Mm(-BODY_RIGHT_MM)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(assets / "body-header.png"), width=Mm(A4_WIDTH_MM))

    footer = section.footer
    footer.is_linked_to_previous = False
    clear_paragraph(footer.paragraphs[0])
    metadata_table = footer.add_table(rows=1, cols=2, width=Mm(CONTENT_WIDTH_MM))
    metadata_table.autofit = False
    set_table_geometry(metadata_table, [CONTENT_WIDTH_MM * 0.72, CONTENT_WIDTH_MM * 0.28])
    for cell in metadata_table.rows[0].cells:
        set_cell_margins(cell, top=20, bottom=20, start=0, end=0)
        tcpr = cell._tc.get_or_add_tcPr()
        borders = tcpr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tcpr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "nil")
            borders.append(node)

    left = metadata_table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source_short = metadata.get('source_hash', '')[:12]
    left_text = (
        f"{metadata['issue_date']}  |  {tr(data, 'version', 'Versión')} "
        f"{metadata['version']}  |  SHA-256 {source_short}"
    )
    set_run_font(left.add_run(left_text), size=7.2, color=COLORS["muted"])

    right = metadata_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_fields(right, data)

    band_paragraph = footer.add_paragraph()
    band_paragraph.paragraph_format.left_indent = Mm(-BODY_LEFT_MM)
    band_paragraph.paragraph_format.right_indent = Mm(-BODY_RIGHT_MM)
    band_paragraph.paragraph_format.space_before = Pt(0)
    band_paragraph.paragraph_format.space_after = Pt(0)
    band_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band_paragraph.add_run().add_picture(str(footer_band), width=Mm(A4_WIDTH_MM))


def add_part_title(document: Document, label: str, title: str):
    paragraph = document.add_paragraph(style="Part Label")
    paragraph.add_run(label.upper())
    paragraph = document.add_paragraph(style="Section Title")
    paragraph.add_run(title)


def add_section_heading(document: Document, number: str, title: str):
    document.add_paragraph(f"{number}. {title}", style="Heading 1")


def add_paragraphs(document: Document, paragraphs: Iterable[str], style: str | None = None):
    for text in paragraphs or []:
        if not str(text).strip():
            continue
        paragraph = document.add_paragraph(style=style)
        paragraph.add_run(str(text))


def add_bullets(document: Document, items: Iterable[str]):
    for item in items or []:
        if not str(item).strip():
            continue
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(str(item))


def create_numbering_instance(document: Document) -> int:
    numbering = document.part.numbering_part.element
    existing = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    num_id = max(existing, default=0) + 1

    style = document.styles["List Number"]
    style_num_pr = style._element.pPr.numPr
    if style_num_pr is None or style_num_pr.numId is None:
        raise ValueError("List Number style has no numbering definition")
    base_num_id = str(style_num_pr.numId.val)
    base_num = numbering.find(
        f".//w:num[@w:numId='{base_num_id}']", namespaces=numbering.nsmap
    )
    if base_num is None:
        raise ValueError("List Number base numbering instance not found")
    abstract = base_num.find(qn("w:abstractNumId"))
    abstract_id = abstract.get(qn("w:val"))

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(num_id))
    new_abstract = OxmlElement("w:abstractNumId")
    new_abstract.set(qn("w:val"), abstract_id)
    new_num.append(new_abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    new_num.append(override)
    numbering.append(new_num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.numPr
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num = num_pr.find(qn("w:numId"))
    if num is None:
        num = OxmlElement("w:numId")
        num_pr.append(num)
    num.set(qn("w:val"), str(num_id))


def add_numbered(document: Document, items: Iterable[str]):
    num_id = create_numbering_instance(document)
    for item in items or []:
        if not str(item).strip():
            continue
        paragraph = document.add_paragraph(style="List Number")
        apply_numbering(paragraph, num_id)
        paragraph.add_run(str(item))


def add_callout(
    document: Document,
    title: str,
    body: str,
    fill: str = COLORS["mint_tint"],
):
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_MM])
    set_table_borders(table, color=COLORS["line"], size=5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=170, end=170)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    set_run_font(paragraph.add_run(title), size=9.5, color=COLORS["navy"], bold=True)
    if body:
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(paragraph.add_run(body), size=9.2, color=COLORS["text"])
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(3)


def add_key_value_table(document: Document, values: dict[str, Any]):
    rows = [[key, value] for key, value in values.items() if str(value).strip()]
    add_table(document, ["Dimensión", "Síntesis"], rows, [43, CONTENT_WIDTH_MM - 43])


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[Any]],
    widths_mm: list[float] | None = None,
    alignments: list[int] | None = None,
):
    if not rows:
        return None
    if widths_mm is None:
        widths_mm = [CONTENT_WIDTH_MM / len(headers)] * len(headers)
    if len(widths_mm) != len(headers):
        raise ValueError("Table widths must match header count")
    scale = CONTENT_WIDTH_MM / sum(widths_mm)
    widths_mm = [value * scale for value in widths_mm]

    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_mm)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, COLORS["navy"])
        set_cell_margins(cell, top=110, bottom=110, start=100, end=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        set_run_font(
            paragraph.add_run(str(text)),
            size=8.2,
            color=COLORS["white"],
            bold=True,
        )

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            if row_index % 2:
                set_cell_shading(cell, COLORS["navy_tint"])
            set_cell_margins(cell, top=105, bottom=105, start=100, end=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if alignments and column_index < len(alignments):
                paragraph.alignment = alignments[column_index]
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(
                paragraph.add_run("" if value is None else str(value)),
                size=8.4,
                color=COLORS["text"],
            )

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(4)
    return table


def rows_from_dicts(items: list[dict[str, Any]], keys: list[str]) -> list[list[Any]]:
    return [[item.get(key, "") for key in keys] for item in items or []]


def _replace_paragraph_text(paragraph, replacements: dict[str, str]):
    current = paragraph.text.strip()
    replacement = replacements.get(current)
    if replacement is None or replacement == current:
        return
    if paragraph.runs:
        paragraph.runs[0].text = replacement
        for run in paragraph.runs[1:]:
            run.text = ''


def translate_document(document: Document, data: dict[str, Any]):
    if data.get('language') != 'english':
        return
    labels = data['labels']
    replacements = {
        'Documento comercial': labels['document_label'],
        'Ficha de propuesta': labels['proposal_sheet'],
        'Contenido': labels['contents'],
        'Parte I': labels['part_one'],
        'Propuesta comercial': labels['commercial_proposal'],
        'PARTE I. PROPUESTA COMERCIAL': f"{labels['part_one'].upper()}. {labels['commercial_proposal'].upper()}",
        'PARTE III. ACEPTACIÓN Y FIRMAS': f"{labels['signatures_part'].upper()}. {labels['signatures_title'].upper()}",
        'ANEXO A. CONTROL DE CAMBIOS': f"{labels['change_part'].upper()}. {labels['change_control'].upper()}",
        '1. Resumen ejecutivo': f"1. {labels['executive_summary']}",
        '2. Antecedentes y entendimiento de la necesidad': f"2. {labels['background']}",
        'Problemas y oportunidades identificados': labels['issues'],
        '3. Objetivos': f"3. {labels['objectives']}",
        '3.1 Objetivo general': f"3.1 {labels['general_objective']}",
        '3.2 Objetivos específicos': f"3.2 {labels['specific_objectives']}",
        '3.3 Indicadores de éxito': f"3.3 {labels['success_indicators']}",
        '4. Alternativas de solución': f"4. {labels['alternatives']}",
        'Alternativa recomendada': labels['recommended_alternative'],
        '5. Alcance, exclusiones y supuestos': f"5. {labels['scope']}",
        '5.1 Alcance incluido': f"5.1 {labels['included_scope']}",
        '5.2 Módulos incluidos': f"5.2 {labels['modules']}",
        '5.3 Fuera de alcance': f"5.3 {labels['exclusions']}",
        '5.4 Supuestos y dependencias': f"5.4 {labels['assumptions']}",
        '6. Metodología y gobierno del proyecto': f"6. {labels['methodology']}",
        '6.1 Prácticas de trabajo': f"6.1 {labels['practices']}",
        '6.2 Gobierno': f"6.2 {labels['governance']}",
        '7. Plan de trabajo, entregables e hitos': f"7. {labels['work_plan']}",
        '7.1 Plan por etapas': f"7.1 {labels['plan_by_stages']}",
        '7.2 Entregables': f"7.2 {labels['deliverables']}",
        '8. Solución técnica e infraestructura': f"8. {labels['technical_solution']}",
        '8.1 Stack tecnológico preliminar': f"8.1 {labels['technology_stack']}",
        '8.2 Requisitos no funcionales': f"8.2 {labels['nonfunctional']}",
        '8.3 Seguridad y continuidad': f"8.3 {labels['security']}",
        '9. Calidad, pruebas y aceptación': f"9. {labels['quality']}",
        '9.1 Estrategia de pruebas': f"9.1 {labels['testing']}",
        '9.2 Procedimiento de aceptación': f"9.2 {labels['acceptance_process']}",
        '9.3 Garantía': f"9.3 {labels['warranty']}",
        '10. Equipo de trabajo': f"10. {labels['team']}",
        '11. Cronograma': f"11. {labels['timeline']}",
        'Condición de inicio': labels['start_condition'],
        '12. Inversión y esquema de pagos': f"12. {labels['investment_payments']}",
        '12.1 Esquema de pagos': f"12.1 {labels['payment_schedule']}",
        '12.2 Condiciones comerciales': f"12.2 {labels['commercial_terms']}",
        '13. Vigencia y aceptación comercial': f"13. {labels['validity_acceptance']}",
        'Condiciones de inicio': labels['start_conditions'],
        'Parte III': labels['signatures_part'],
        'Aceptación y firmas': labels['signatures_title'],
        'Datos de aceptación': labels['acceptance_data'],
        'Firmas': labels['signatures'],
        'Anexo A': labels['change_part'],
        'Control de cambios': labels['change_control'],
        'Procedimiento': labels['procedure'],
        'Formulario base': labels['base_form'],
    }
    spanish_by_key = {
        'field': 'Campo', 'detail': 'Detalle', 'client': 'Cliente',
        'project': 'Proyecto', 'version': 'Versión', 'issue_date': 'Fecha de emisión',
        'validity': 'Vigencia', 'status': 'Estado', 'dimension': 'Dimensión',
        'summary': 'Síntesis', 'indicator': 'Indicador', 'target': 'Objetivo',
        'verification': 'Verificación', 'alternative': 'Alternativa',
        'duration': 'Duración', 'investment': 'Inversión', 'limits': 'Límites',
        'module': 'Módulo', 'capabilities': 'Capacidades', 'limit': 'Límite',
        'role': 'Rol', 'responsibility': 'Responsabilidad', 'cadence': 'Frecuencia',
        'stage': 'Etapa', 'activities': 'Actividades', 'deliverables': 'Entregables',
        'milestone': 'Hito', 'code': 'Código', 'deliverable': 'Entregable',
        'format': 'Formato', 'minimum_content': 'Contenido mínimo',
        'acceptance': 'Aceptación', 'layer': 'Capa', 'technology': 'Tecnología',
        'notes': 'Observaciones', 'team': 'Equipo de trabajo',
        'dedication': 'Dedicación estimada', 'timing': 'Momento',
        'validates': 'Valida', 'dependencies': 'Dependencias',
        'percentage': 'Porcentaje', 'verifiable_condition': 'Condición verificable',
        'organization': 'Organización', 'name': 'Nombre y apellido',
        'personal_id': 'DNI / identificación', 'email': 'Correo electrónico',
        'date': 'Fecha', 'signature': 'Firma', 'request_number': 'Número de solicitud',
        'requester': 'Solicitante', 'description': 'Descripción',
        'scope_impact': 'Impacto en alcance', 'schedule_impact': 'Impacto en cronograma',
        'fees_impact': 'Impacto en honorarios', 'resolution': 'Resolución',
        'pending': 'Pendiente',
    }
    for key, spanish in spanish_by_key.items():
        replacements[spanish] = labels[key]

    containers = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for section in document.sections:
        for area in (section.header, section.footer):
            containers.extend(area.paragraphs)
            for table in area.tables:
                for row in table.rows:
                    for cell in row.cells:
                        containers.extend(cell.paragraphs)
    for paragraph in containers:
        _replace_paragraph_text(paragraph, replacements)


def add_metadata_page(document: Document, data: dict[str, Any], contents: list[str]):
    metadata = data["metadata"]
    add_part_title(
        document,
        tr(data, 'document_label', 'Documento comercial'),
        tr(data, 'proposal_sheet', 'Ficha de propuesta'),
    )
    add_table(
        document,
        [tr(data, 'field', 'Campo'), tr(data, 'detail', 'Detalle')],
        [
            [tr(data, 'client', 'Cliente'), metadata["client"]],
            [tr(data, 'project', 'Proyecto'), metadata["project_name"]],
            [tr(data, 'version', 'Versión'), metadata["version"]],
            [tr(data, 'issue_date', 'Fecha de emisión'), metadata["issue_date"]],
            [tr(data, 'validity', 'Vigencia'), metadata["valid_until"]],
            [tr(data, 'status', 'Estado'), metadata.get("status", "")],
            [tr(data, 'source_version', 'Versión de fuente'), metadata.get('source_version', '')],
            [tr(data, 'source_hash', 'SHA-256 de fuente'), metadata.get('source_hash', '')],
            [tr(data, 'generated_at', 'Generado'), metadata.get('generated_at', '')],
        ],
        [42, CONTENT_WIDTH_MM - 42],
    )
    document.add_paragraph(tr(data, 'contents', 'Contenido'), style="Heading 1")
    for item in contents:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(item), size=9, color=COLORS["text"])
    document.add_page_break()


def add_executive_summary(document: Document, data: dict[str, Any]):
    section = data["executive_summary"]
    add_section_heading(document, "1", "Resumen ejecutivo")
    add_paragraphs(document, section.get("paragraphs", []), style="Lead")
    add_key_value_table(document, section.get("highlights", {}))


def add_context(document: Document, data: dict[str, Any]):
    section = data["context"]
    add_section_heading(document, "2", "Antecedentes y entendimiento de la necesidad")
    add_paragraphs(document, section.get("paragraphs", []))
    if section.get("pain_points"):
        document.add_paragraph("Problemas y oportunidades identificados", style="Heading 2")
        add_bullets(document, section["pain_points"])


def add_objectives(document: Document, data: dict[str, Any]):
    section = data["objectives"]
    add_section_heading(document, "3", "Objetivos")
    document.add_paragraph("3.1 Objetivo general", style="Heading 2")
    add_paragraphs(document, [section.get("general", "")])
    document.add_paragraph("3.2 Objetivos específicos", style="Heading 2")
    add_bullets(document, section.get("specific", []))
    indicators = section.get("success_indicators", [])
    if indicators:
        document.add_paragraph("3.3 Indicadores de éxito", style="Heading 2")
        add_table(
            document,
            ["Indicador", "Objetivo", "Verificación"],
            rows_from_dicts(indicators, ["indicator", "target", "verification"]),
            [48, 32, 94],
        )


def add_alternatives(document: Document, data: dict[str, Any]):
    alternatives = data.get("alternatives", [])
    if not alternatives:
        return
    add_section_heading(document, "4", "Alternativas de solución")
    add_table(
        document,
        ["Alternativa", "Síntesis", "Duración", "Inversión", "Límites"],
        rows_from_dicts(
            alternatives, ["name", "summary", "duration", "investment", "limits"]
        ),
        [34, 61, 22, 25, 32],
    )
    recommendation = data.get("recommended_alternative")
    if recommendation:
        add_callout(
            document,
            "Alternativa recomendada",
            recommendation,
            fill=COLORS["mint_tint"],
        )


def add_scope(document: Document, data: dict[str, Any]):
    section = data["scope"]
    add_section_heading(document, "5", "Alcance, exclusiones y supuestos")
    document.add_paragraph("5.1 Alcance incluido", style="Heading 2")
    add_bullets(document, section.get("included", []))
    modules = section.get("modules", [])
    if modules:
        document.add_paragraph("5.2 Módulos incluidos", style="Heading 2")
        add_table(
            document,
            ["Módulo", "Capacidades", "Límite"],
            rows_from_dicts(modules, ["module", "capabilities", "limit"]),
            [38, 90, 46],
        )
    document.add_paragraph("5.3 Fuera de alcance", style="Heading 2")
    add_bullets(document, section.get("exclusions", []))
    document.add_paragraph("5.4 Supuestos y dependencias", style="Heading 2")
    add_bullets(document, section.get("assumptions", []))


def add_methodology(document: Document, data: dict[str, Any]):
    section = data["methodology"]
    add_section_heading(document, "6", "Metodología y gobierno del proyecto")
    add_paragraphs(document, [section.get("approach", "")])
    document.add_paragraph("6.1 Prácticas de trabajo", style="Heading 2")
    add_bullets(document, section.get("practices", []))
    governance = section.get("governance", [])
    if governance:
        document.add_paragraph("6.2 Gobierno", style="Heading 2")
        add_table(
            document,
            ["Rol", "Responsabilidad", "Frecuencia"],
            rows_from_dicts(governance, ["role", "responsibility", "cadence"]),
            [45, 91, 38],
        )


def add_work_plan(document: Document, data: dict[str, Any]):
    section = data["work_plan"]
    add_section_heading(document, "7", "Plan de trabajo, entregables e hitos")
    stages = section.get("stages", [])
    if stages:
        document.add_paragraph("7.1 Plan por etapas", style="Heading 2")
        add_table(
            document,
            ["Etapa", "Actividades", "Entregables", "Duración", "Hito"],
            rows_from_dicts(
                stages,
                ["stage", "activities", "deliverables", "duration", "milestone"],
            ),
            [30, 48, 40, 22, 34],
        )
    deliverables = section.get("deliverables", [])
    if deliverables:
        document.add_paragraph("7.2 Entregables", style="Heading 2")
        add_table(
            document,
            ["Código", "Entregable", "Formato", "Contenido mínimo", "Aceptación"],
            rows_from_dicts(
                deliverables,
                ["code", "name", "format", "minimum_content", "acceptance"],
            ),
            [18, 39, 23, 48, 46],
        )


def add_technical_solution(document: Document, data: dict[str, Any]):
    section = data.get("technical_solution", {})
    add_section_heading(
        document, "8", tr(data, 'technical_solution', 'Solución técnica e infraestructura')
    )
    add_paragraphs(document, section.get("architecture", []))
    stack = section.get("stack", [])
    if stack:
        document.add_paragraph("8.1 Stack tecnológico preliminar", style="Heading 2")
        add_table(
            document,
            ["Capa", "Tecnología", "Observaciones"],
            rows_from_dicts(stack, ["layer", "technology", "notes"]),
            [38, 48, 88],
        )
    if section.get("nonfunctional"):
        document.add_paragraph("8.2 Requisitos no funcionales", style="Heading 2")
        add_bullets(document, section["nonfunctional"])
    if section.get("security"):
        document.add_paragraph("8.3 Seguridad y continuidad", style="Heading 2")
        add_bullets(document, section["security"])


def add_quality(document: Document, data: dict[str, Any]):
    section = data["quality_acceptance"]
    add_section_heading(document, "9", "Calidad, pruebas y aceptación")
    document.add_paragraph("9.1 Estrategia de pruebas", style="Heading 2")
    add_bullets(document, section.get("testing", []))
    document.add_paragraph("9.2 Procedimiento de aceptación", style="Heading 2")
    add_numbered(document, section.get("acceptance_steps", []))
    if section.get("warranty"):
        document.add_paragraph("9.3 Garantía", style="Heading 2")
        add_paragraphs(document, [section["warranty"]])


def add_team(document: Document, data: dict[str, Any]):
    add_section_heading(document, "10", "Equipo de trabajo")
    add_table(
        document,
        ["Rol", "Responsabilidad", "Dedicación estimada"],
        rows_from_dicts(data.get("team", []), ["role", "responsibility", "dedication"]),
        [44, 88, 42],
    )


def add_timeline(document: Document, data: dict[str, Any]):
    section = data["timeline"]
    add_section_heading(document, "11", "Cronograma")
    add_callout(
        document,
        "Condición de inicio",
        section.get("start_condition", ""),
        fill=COLORS["navy_tint"],
    )
    add_table(
        document,
        ["Hito", "Momento", "Valida", "Dependencias"],
        rows_from_dicts(
            section.get("milestones", []),
            ["milestone", "timing", "validator", "dependencies"],
        ),
        [46, 28, 42, 58],
    )


def add_commercial(document: Document, data: dict[str, Any]):
    section = data["commercial"]
    add_section_heading(document, "12", "Inversión y esquema de pagos")
    add_table(
        document,
        ["Alternativa", "Inversión", "Duración"],
        rows_from_dicts(section.get("options", []), ["alternative", "amount", "duration"]),
        [86, 44, 44],
    )

    payment_schedule = section.get("payment_schedule", [])
    if payment_schedule:
        document.add_paragraph("12.1 Esquema de pagos", style="Heading 2")
        add_table(
            document,
            ["Hito", "Porcentaje", "Condición verificable"],
            rows_from_dicts(payment_schedule, ["milestone", "percentage", "condition"]),
            [58, 28, 88],
        )
    if section.get("commercial_terms"):
        document.add_paragraph("12.2 Condiciones comerciales", style="Heading 2")
        add_bullets(document, section["commercial_terms"])


def add_validity(document: Document, data: dict[str, Any]):
    section = data["validity_acceptance"]
    add_section_heading(document, "13", "Vigencia y aceptación comercial")
    add_paragraphs(document, section.get("paragraphs", []))
    if section.get("conditions"):
        document.add_paragraph("Condiciones de inicio", style="Heading 2")
        add_bullets(document, section["conditions"])


def add_legal_paragraphs(document: Document, paragraphs: Iterable[str]):
    for text in paragraphs or []:
        if not str(text).strip():
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2.5)
        paragraph.paragraph_format.line_spacing = 1.0
        set_run_font(paragraph.add_run(str(text)), size=9, color=COLORS["text"])


def add_legal_bullets(document: Document, items: Iterable[str]):
    for item in items or []:
        if not str(item).strip():
            continue
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.paragraph_format.line_spacing = 1.0
        set_run_font(paragraph.add_run(str(item)), size=9, color=COLORS["text"])


def add_legal_numbered(document: Document, items: Iterable[str]):
    items = list(items or [])
    if not items:
        return
    num_id = create_numbering_instance(document)
    for item in items:
        if not str(item).strip():
            continue
        paragraph = document.add_paragraph(style="List Number")
        apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.paragraph_format.line_spacing = 1.0
        set_run_font(paragraph.add_run(str(item)), size=9, color=COLORS["text"])
def load_general_terms(assets: Path, asset_name: str) -> dict[str, Any]:
    path = assets / asset_name
    return json.loads(path.read_text(encoding="utf-8"))


def add_terms(document: Document, data: dict[str, Any], assets: Path):
    document.add_page_break()
    add_part_title(document, "Parte II", "Términos y condiciones")
    asset_name = data.get('options', {}).get('general_terms_asset', 'general-terms.es.json')
    general_terms = load_general_terms(assets, asset_name)
    add_legal_paragraphs(document, general_terms.get("intro", []))
    clauses = list(general_terms.get("clauses", []))
    additional = data.get("terms_additional")
    if additional is None:
        additional = data.get("terms", [])
    clauses.extend(additional or [])

    for index, clause in enumerate(clauses, 1):
        heading = document.add_paragraph(f"{index}. {clause.get('title', '')}", style="Heading 2")
        heading.paragraph_format.space_before = Pt(5)
        heading.paragraph_format.space_after = Pt(2)
        add_legal_paragraphs(document, clause.get("paragraphs", []))
        add_legal_numbered(document, clause.get("numbered", []))
        add_legal_bullets(document, clause.get("bullets", []))
        for subsection_index, subsection in enumerate(clause.get("subsections", []), 1):
            subheading = document.add_paragraph(
                f"{index}.{subsection_index} {subsection.get('title', '')}",
                style="Heading 3",
            )
            subheading.paragraph_format.space_before = Pt(3.5)
            subheading.paragraph_format.space_after = Pt(1.5)
            add_legal_paragraphs(document, subsection.get("paragraphs", []))
            add_legal_numbered(document, subsection.get("numbered", []))
            add_legal_bullets(document, subsection.get("bullets", []))


def add_signature_field(
    cell, label: str, tag: str, value: str = "", signature=False,
    empty_text: str = "Haga clic o toque aquí para completar",
    signature_text: str = "Insertar firma electrónica o digital",
):
    label_paragraph = cell.add_paragraph()
    label_paragraph.paragraph_format.space_before = Pt(8 if signature else 4)
    label_paragraph.paragraph_format.space_after = Pt(1)
    set_run_font(
        label_paragraph.add_run(label),
        size=8.2,
        color=COLORS["text"],
        bold=True,
    )
    field_paragraph = cell.add_paragraph()
    field_paragraph.paragraph_format.space_after = Pt(7 if signature else 3)
    if signature:
        field_paragraph.paragraph_format.space_before = Pt(5)
        field_paragraph.paragraph_format.space_after = Pt(16)
    add_text_content_control(
        field_paragraph,
        tag,
        label,
        value or (signature_text if signature else empty_text),
        size=8.3,
    )


def add_digital_acceptance_table(document: Document, data: dict[str, Any]):
    metadata = data["metadata"]
    signatures = data.get("signatures", {})
    client = signatures.get("client", {})
    rows = [
        (
            tr(data, 'selected_scope', 'Alternativa o alcance seleccionado'),
            "acceptance.selected_scope",
            signatures.get("selected_alternative", ""),
        ),
        (
            tr(data, 'client', 'Razón social del cliente'),
            "acceptance.client.organization",
            client.get("organization", metadata.get("client", "")),
        ),
        (
            tr(data, 'tax_id', 'CUIT / identificación fiscal'),
            "acceptance.client.tax_id",
            client.get("tax_id", ""),
        ),
        (
            tr(data, 'authorized_representative', 'Representante autorizado'),
            "acceptance.client.representative",
            client.get("name", ""),
        ),
        (
            tr(data, 'email', 'Correo electrónico'),
            "acceptance.client.email",
            client.get("email", ""),
        ),
        (
            tr(data, 'acceptance_date', 'Fecha de aceptación'),
            "acceptance.date",
            signatures.get("acceptance_date", ""),
        ),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    set_table_geometry(table, [58, CONTENT_WIDTH_MM - 58])
    set_table_borders(table)
    for row_index, (label, tag, value) in enumerate(rows):
        label_cell = table.cell(row_index, 0)
        value_cell = table.cell(row_index, 1)
        set_cell_shading(label_cell, COLORS["navy_tint"])
        for cell in (label_cell, value_cell):
            set_cell_margins(cell, top=120, bottom=120, start=130, end=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_paragraph = label_cell.paragraphs[0]
        set_run_font(
            label_paragraph.add_run(label),
            size=8.3,
            color=COLORS["navy"],
            bold=True,
        )
        add_text_content_control(
            value_cell.paragraphs[0],
            tag,
            label,
            value or tr(data, 'click_to_fill', 'Haga clic o toque aquí para completar'),
            size=8.3,
        )
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_signatures(document: Document, data: dict[str, Any]):
    document.add_page_break()
    add_part_title(
        document,
        tr(data, 'signatures_part', 'Parte III'),
        tr(data, 'signatures_title', 'Aceptación y firmas'),
    )
    metadata = data["metadata"]
    add_paragraphs(
        document,
        [
            tr(data, 'signatures_intro_1', 'Complete los campos para identificar la propuesta y versión aceptadas.'),
            tr(data, 'signatures_intro_2', 'Los campos son controles editables de Word.'),
        ],
    )
    document.add_paragraph(tr(data, 'acceptance_data', 'Datos de aceptación'), style="Heading 2")
    add_digital_acceptance_table(document, data)

    document.add_paragraph(tr(data, 'signatures', 'Firmas'), style="Heading 2")
    signatures = data.get("signatures", {})
    client = signatures.get("client", {})
    quasar = signatures.get("provider", signatures.get("quasar", {}))
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [CONTENT_WIDTH_MM / 2, CONTENT_WIDTH_MM / 2])
    set_table_borders(table)
    parties = (
        (tr(data, 'client_party', 'POR EL CLIENTE'), "signature.client", client, metadata.get("client", "")),
        (tr(data, 'quasar_party', 'POR INGENIERÍA QUASAR SRL'), "signature.quasar", quasar, LEGAL_ENTITY),
    )
    for index, (label, tag_prefix, party, default_organization) in enumerate(parties):
        cell = table.cell(0, index)
        set_cell_shading(cell, COLORS["navy_tint"])
        set_cell_margins(cell, top=160, bottom=160, start=160, end=160)
        paragraph = cell.paragraphs[0]
        set_run_font(
            paragraph.add_run(label),
            size=9,
            color=COLORS["navy"],
            bold=True,
        )
        fields = (
            ("Organización", "organization", party.get("organization", default_organization), False),
            ("Nombre y apellido", "name", party.get("name", ""), False),
            ("DNI / identificación", "id", party.get("id", ""), False),
            ("Cargo", "role", party.get("role", ""), False),
            ("Correo electrónico", "email", party.get("email", ""), False),
            ("Fecha", "date", party.get("date", ""), False),
            ("Firma", "signature", "", True),
        )
        for field_label, field_tag, value, is_signature in fields:
            add_signature_field(
                cell,
                field_label,
                f"{tag_prefix}.{field_tag}",
                value,
                signature=is_signature,
                empty_text=tr(data, 'click_to_fill', 'Haga clic o toque aquí para completar'),
                signature_text=tr(data, 'insert_signature', 'Insertar firma electrónica o digital'),
            )

def add_change_control(document: Document, data: dict[str, Any]):
    section = data.get("change_control", {})
    document.add_page_break()
    add_part_title(
        document,
        tr(data, 'change_part', 'Anexo A'),
        tr(data, 'change_control', 'Control de cambios'),
    )
    add_paragraphs(document, section.get('paragraphs', [section.get('description', '')]))
    steps = section.get('steps', section.get('procedure', []))
    if steps:
        document.add_paragraph(tr(data, 'procedure', 'Procedimiento'), style="Heading 2")
        add_numbered(document, steps)
    fields = section.get('form_fields', section.get('fields', []))
    rows = []
    if isinstance(fields, dict):
        rows = [[key, value] for key, value in fields.items()]
    elif isinstance(fields, list):
        rows = [
            [item.get('label', item.get('name', '')), item.get('value', '')]
            for item in fields if isinstance(item, dict)
        ]
    if rows:
        document.add_paragraph(tr(data, 'base_form', 'Formulario base'), style="Heading 2")
        add_table(
            document,
            [tr(data, 'field', 'Campo'), tr(data, 'detail', 'Contenido')],
            rows,
            [48, CONTENT_WIDTH_MM - 48],
        )


def contents_for(data: dict[str, Any]) -> list[str]:
    options = data.get("options", {})
    items = [
        "PARTE I. PROPUESTA COMERCIAL",
        "1. Resumen ejecutivo",
        "2. Antecedentes y entendimiento de la necesidad",
        "3. Objetivos",
    ]
    if options.get("include_alternatives", bool(data.get("alternatives"))):
        items.append("4. Alternativas de solución")
    items.extend(
        [
            "5. Alcance, exclusiones y supuestos",
            "6. Metodología y gobierno del proyecto",
            "7. Plan de trabajo, entregables e hitos",
        ]
    )
    if options.get("include_technical_solution", True):
        items.append(f"8. {tr(data, 'technical_solution', 'Solución técnica e infraestructura')}")
    items.extend(
        [
            "9. Calidad, pruebas y aceptación",
            "10. Equipo de trabajo",
            "11. Cronograma",
            "12. Inversión y esquema de pagos",
            "13. Vigencia y aceptación comercial",
        ]
    )
    if options.get("include_terms", True):
        items.append("PARTE II. TÉRMINOS Y CONDICIONES")
    if options.get("include_signatures", True):
        items.append("PARTE III. ACEPTACIÓN Y FIRMAS")
    if options.get("include_change_control", True):
        items.append("ANEXO A. CONTROL DE CAMBIOS")
    return items


def build_document(data: dict[str, Any], output: Path, assets: Path, keep_cover=False):
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="quasar-proposal-") as tmp:
        tmp_path = Path(tmp)
        cover_path = tmp_path / "cover.png"
        footer_band = tmp_path / "footer-band.png"
        build_cover_image(data, assets, cover_path)
        build_footer_band(data, assets, footer_band)

        document = Document()
        configure_styles(document)
        set_update_fields(document)
        metadata = data["metadata"]
        document.core_properties.title = (
            f"Propuesta comercial - {metadata['project_name']}"
        )
        document.core_properties.subject = metadata.get("project_subtitle", "")
        document.core_properties.author = LEGAL_ENTITY
        source_hash = metadata.get('source_hash', '')
        source_version = metadata.get('source_version', metadata.get('version', ''))
        document.core_properties.keywords = (
            f"Quasar; derived; proposal-source-v{source_version}; sha256:{source_hash}"
        )
        document.core_properties.comments = (
            f"creation_mode=derived; semantic_authority=none; "
            f"source_version={source_version}; source_sha256={source_hash}; "
            f"generated_by=commercial-proposal-document; "
            f"generated_at={metadata.get('generated_at', '')}; do_not_edit=true"
        )

        cover_section = document.sections[0]
        configure_cover_section(cover_section)
        cover_paragraph = document.add_paragraph()
        cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_paragraph.paragraph_format.space_before = Pt(0)
        cover_paragraph.paragraph_format.space_after = Pt(0)
        cover_paragraph.add_run().add_picture(
            str(cover_path), width=Mm(208.5), height=Mm(294.8)
        )

        body_section = document.add_section(WD_SECTION.NEW_PAGE)
        configure_body_section(body_section, assets, footer_band, data)

        add_metadata_page(document, data, contents_for(data))
        add_part_title(document, "Parte I", "Propuesta comercial")
        add_executive_summary(document, data)
        add_context(document, data)
        add_objectives(document, data)

        options = data.get("options", {})
        if options.get("include_alternatives", bool(data.get("alternatives"))):
            add_alternatives(document, data)
        add_scope(document, data)
        add_methodology(document, data)
        add_work_plan(document, data)
        if options.get("include_technical_solution", True):
            add_technical_solution(document, data)
        add_quality(document, data)
        add_team(document, data)
        add_timeline(document, data)
        add_commercial(document, data)
        add_validity(document, data)
        if options.get("include_terms", True):
            add_terms(document, data, assets)
        if options.get("include_signatures", True):
            add_signatures(document, data)
        if options.get("include_change_control", True):
            add_change_control(document, data)

        translate_document(document, data)
        document.save(output)
        if keep_cover:
            cover_copy = output.with_name(output.stem + "-cover.png")
            cover_copy.write_bytes(cover_path.read_bytes())
