from __future__ import annotations

import argparse
import json
import re
import unicodedata
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from document_model import (
    as_dict, build_render_model, expected_facts, load_yaml, model_completeness,
    parse_mapping, schema_errors, sha256_file, validate_source_and_mapping,
)


PLACEHOLDERS = [
    re.compile(r'\{\{[^}]+\}\}'), re.compile(r'<[^<>]{2,80}>'),
    re.compile(r'\b(?:PENDIENTE|PENDING)\s*:', re.IGNORECASE),
    re.compile(r'\b(?:TBD|XXX|TO[ -]?DO)\b', re.IGNORECASE),
]


def normalize(value: str) -> str:
    value = unicodedata.normalize('NFKC', value or '').casefold()
    return ' '.join(value.split())


def collect_docx_text(document: Any) -> str:
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    for section in document.sections:
        for area in (section.header, section.footer):
            chunks.extend(paragraph.text for paragraph in area.paragraphs)
            for table in area.tables:
                for row in table.rows:
                    for cell in row.cells:
                        chunks.append(cell.text)
    return '\n'.join(chunks)


def required_headings(model: dict[str, Any]) -> list[str]:
    labels = model['labels']
    headings = [
        f"1. {labels['executive_summary']}", f"2. {labels['background']}",
        f"3. {labels['objectives']}", f"5. {labels['scope']}",
        f"6. {labels['methodology']}", f"7. {labels['work_plan']}",
        f"9. {labels['quality']}", f"10. {labels['team']}",
        f"11. {labels['timeline']}", f"12. {labels['investment_payments']}",
        f"13. {labels['validity_acceptance']}",
    ]
    options = as_dict(model.get('options'))
    if options.get('include_alternatives'):
        headings.append(f"4. {labels['alternatives']}")
    if options.get('include_technical_solution'):
        headings.append(f"8. {labels['technical_solution']}")
    return headings


def missing_facts(haystack: str, facts: list[str]) -> list[str]:
    normalized = normalize(haystack)
    return [
        fact for fact in facts
        if len(normalize(fact)) >= 3 and normalize(fact) not in normalized
    ]


def validate_docx(
    path: Path, model: dict[str, Any], assets: Path
) -> tuple[list[str], list[str], dict[str, Any]]:
    try:
        from docx import Document
        from document_builder import CONTACT_EMAIL, WEBSITE
    except ModuleNotFoundError as exc:
        dependency = {'docx': 'python-docx', 'PIL': 'Pillow'}.get(
            exc.name or '', exc.name or 'document runtime'
        )
        return [
            f'Missing dependency: {dependency} is required for DOCX validation'
        ], [], {}

    errors: list[str] = []
    warnings: list[str] = []
    document = Document(path)
    docx_text = collect_docx_text(document)
    mode = as_dict(model.get('options')).get('mode', 'draft')
    metadata = as_dict(model.get('metadata'))

    if len(document.sections) < 2:
        errors.append('DOCX has no separate cover and interior sections')
    for index, section in enumerate(document.sections, 1):
        width = round(section.page_width.mm, 1)
        height = round(section.page_height.mm, 1)
        if abs(width - 210) > 1 or abs(height - 297) > 1:
            errors.append(f'Section {index} is not A4: {width} x {height} mm')
    for heading in required_headings(model):
        if heading not in docx_text:
            errors.append(f'Missing required section: {heading}')

    styles = {style.name for style in document.styles}
    for required in ('Heading 1', 'Heading 2', 'Part Label', 'Section Title'):
        if required not in styles:
            errors.append(f'Missing Word style: {required}')
    heading1 = document.styles['Heading 1']
    if heading1.font.color.rgb is None or str(heading1.font.color.rgb) != '27367E':
        errors.append('Heading 1 does not use Quasar navy #27367E')
    if not document.tables:
        errors.append('DOCX contains no structured tables')

    comments = document.core_properties.comments or ''
    source_hash = metadata.get('source_hash', '')
    source_version = metadata.get('source_version', '')
    if source_hash not in comments or f'source_version={source_version}' not in comments:
        errors.append('DOCX core properties do not contain source version and SHA-256')
    if source_hash not in docx_text:
        errors.append('DOCX does not show the complete source SHA-256')

    placeholder_hits = []
    for pattern in PLACEHOLDERS:
        placeholder_hits.extend(pattern.findall(docx_text))
    if placeholder_hits:
        message = f'DOCX contains placeholders: {sorted(set(placeholder_hits))[:12]}'
        (errors if mode == 'issued' else warnings).append(message)

    with zipfile.ZipFile(path) as archive:
        media = [name for name in archive.namelist() if name.startswith('word/media/')]
        if len(media) < 3:
            errors.append('DOCX is missing cover, logo, header, or footer media')
        document_xml = archive.read('word/document.xml')
        control_count = document_xml.count(b'<w:sdt>')
        if as_dict(model.get('options')).get('include_signatures') and control_count < 20:
            errors.append(f'Acceptance and signatures have too few editable controls: {control_count}')

    if as_dict(model.get('options')).get('include_terms'):
        asset_name = as_dict(model.get('options')).get(
            'general_terms_asset', 'general-terms.es.json'
        )
        terms = json.loads((assets / asset_name).read_text(encoding='utf-8'))
        for clause in terms.get('clauses', []):
            if clause.get('title') not in docx_text:
                errors.append(f"Missing authorized general clause: {clause.get('title')}")

    lowered = docx_text.casefold()
    if 'quasaranalytic' in lowered or 'quasar analytic' in lowered:
        errors.append('Legacy Quasar Analytic branding was detected')
    if CONTACT_EMAIL != 'info@quasartech.xyz' or WEBSITE != 'quasartech.xyz':
        errors.append('Renderer does not use the approved Quasar contact data')

    facts = expected_facts(model)
    missing = missing_facts(docx_text, facts)
    for fact in missing:
        errors.append(f'DOCX diverges from the canonical proposal source; missing fact: {fact}')
    return errors, warnings, {
        'sections': len(document.sections), 'tables': len(document.tables),
        'editable_controls': control_count, 'canonical_facts_checked': len(facts),
        'canonical_facts_missing': missing,
    }


def validate_pdf(
    path: Path, model: dict[str, Any]
) -> tuple[list[str], list[str], dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ['pypdf is unavailable; PDF validation cannot run'], [], {}
    errors: list[str] = []
    reader = PdfReader(str(path))
    pdf_text = '\n'.join(page.extract_text() or '' for page in reader.pages)
    if not reader.pages:
        errors.append('PDF has no pages')
    if not pdf_text.strip():
        errors.append('PDF text is not extractable')
    facts = expected_facts(model)
    missing = missing_facts(pdf_text, facts)
    for fact in missing:
        errors.append(f'PDF diverges from the canonical proposal source; missing fact: {fact}')
    return errors, [], {
        'pages': len(reader.pages), 'canonical_facts_checked': len(facts),
        'canonical_facts_missing': missing,
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description='Validate a canonical Quasar DOCX/PDF commercial package.'
    )
    parser.add_argument('--source', required=True, type=Path)
    parser.add_argument('--mapping', required=True, type=Path)
    parser.add_argument('--docx', required=True, type=Path)
    parser.add_argument('--pdf', type=Path)
    parser.add_argument('--report', required=True, type=Path)
    parser.add_argument(
        '--visual-review', choices=['passed', 'failed', 'not-run'], default='not-run'
    )
    parser.add_argument('--source-schema', type=Path)
    parser.add_argument('--mapping-schema', type=Path)
    args = parser.parse_args()

    skill_root = Path(__file__).resolve().parents[1]
    source_path = args.source.resolve()
    mapping_path = args.mapping.resolve()
    source_schema = args.source_schema or (
        skill_root.parent / 'q-proposal-design' /
        'references' / '02-proposal-source.schema.yaml'
    )
    mapping_schema = args.mapping_schema or (
        skill_root / 'references' / '04-document-mapping.schema.yaml'
    )
    report_schema = skill_root / 'references' / '04-validation-report.schema.json'

    source = load_yaml(source_path)
    mapping = parse_mapping(mapping_path)
    errors, warnings = validate_source_and_mapping(
        source, mapping, source_path, source_schema.resolve(), mapping_schema.resolve()
    )
    model = build_render_model(source, mapping, source_path)
    completeness_errors, completeness_warnings = model_completeness(model)
    errors.extend(completeness_errors)
    warnings.extend(completeness_warnings)

    docx_errors, docx_warnings, docx_checks = validate_docx(
        args.docx.resolve(), model, skill_root / 'assets'
    )
    errors.extend(docx_errors)
    warnings.extend(docx_warnings)
    checks: dict[str, Any] = {
        'source_schema': (
            'Passed' if not schema_errors(source, source_schema.resolve()) else 'Failed'
        ),
        'mapping_schema': (
            'Passed' if not schema_errors(mapping, mapping_schema.resolve()) else 'Failed'
        ),
        'docx': docx_checks,
        'visual_review': args.visual_review,
    }
    outputs: dict[str, Any] = {
        'docx': {
            'path': str(args.docx.resolve()),
            'sha256': sha256_file(args.docx.resolve()),
        },
        'pdf': None,
    }

    pdf_required = bool(as_dict(mapping.get('document')).get('pdf_required'))
    if args.pdf:
        pdf_errors, pdf_warnings, pdf_checks = validate_pdf(args.pdf.resolve(), model)
        errors.extend(pdf_errors)
        warnings.extend(pdf_warnings)
        checks['pdf'] = pdf_checks
        outputs['pdf'] = {
            'path': str(args.pdf.resolve()),
            'sha256': sha256_file(args.pdf.resolve()),
        }
    elif pdf_required:
        errors.append('PDF is required by the document mapping but was not provided')

    mode = as_dict(model.get('options')).get('mode', 'draft')
    if args.visual_review == 'failed':
        errors.append('Visual review failed')
    elif args.visual_review != 'passed':
        (errors if mode == 'issued' else warnings).append(
            'Visual review has not been completed'
        )

    generated_at = datetime.now(timezone.utc).isoformat()
    status = 'Failed' if errors else ('Passed with warnings' if warnings else 'Passed')
    source_hash = sha256_file(source_path)
    mapping_hash = sha256_file(mapping_path)
    report = {
        '_artifact': {
            'creation_mode': 'derived', 'semantic_authority': 'none',
            'generated_from': [
                {
                    'artifact_id': as_dict(source.get('artifact')).get('artifact_id'),
                    'version': as_dict(source.get('proposal')).get('version'),
                    'hash': source_hash,
                },
                {
                    'artifact_id': as_dict(mapping.get('artifact')).get('artifact_id'),
                    'version': as_dict(mapping.get('artifact')).get('version'),
                    'hash': mapping_hash,
                },
            ],
            'generated_by': 'q-proposal-document',
            'generated_at': generated_at, 'do_not_edit': True,
        },
        'schema_version': '1.0', 'status': status, 'ok': not errors,
        'source': {
            'path': str(source_path),
            'artifact_id': as_dict(source.get('artifact')).get('artifact_id'),
            'version': as_dict(source.get('proposal')).get('version'),
            'sha256': source_hash,
        },
        'mapping': {
            'path': str(mapping_path),
            'version': as_dict(mapping.get('artifact')).get('version'),
            'sha256': mapping_hash,
        },
        'outputs': outputs, 'checks': checks, 'errors': errors, 'warnings': warnings,
    }
    report_schema_errors = schema_errors(report, report_schema)
    if report_schema_errors:
        report['errors'].extend(
            f'Validation report schema: {message}' for message in report_schema_errors
        )
        report['ok'] = False
        report['status'] = 'Failed'

    args.report.parent.mkdir(parents=True, exist_ok=True)
    rendered = json.dumps(report, ensure_ascii=False, indent=2)
    args.report.write_text(rendered, encoding='utf-8')
    print(rendered)
    raise SystemExit(0 if report['ok'] else 1)


if __name__ == '__main__':
    main()
